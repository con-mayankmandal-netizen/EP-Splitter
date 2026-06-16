"""
DOCX Episode Splitter — FastAPI Backend
Detects episodes via configurable regex patterns, splits DOCX files,
and returns a ZIP archive for download.
"""

from __future__ import annotations

import os
import re
import uuid
import shutil
import zipfile
import tempfile
from pathlib import Path
from typing import Optional

import uvicorn
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document

# ─────────────────────────────────────────────────────────────
# App setup
# ─────────────────────────────────────────────────────────────
app = FastAPI(
    title="DOCX Episode Splitter",
    description="Upload a multi-episode DOCX → detect → split → download ZIP",
    version="1.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = Path(tempfile.gettempdir()) / "ep_splitter"
UPLOAD_DIR.mkdir(exist_ok=True)

# In-memory job store  {job_id → dict}
JOBS: dict[str, dict] = {}


# ─────────────────────────────────────────────────────────────
# Episode detection — add new patterns here, order = priority
# ─────────────────────────────────────────────────────────────
EPISODE_PATTERNS: list[re.Pattern] = [
    re.compile(r"^Ep\s+(\d+)\s*[-–—]\s*(.+)$", re.IGNORECASE),
    re.compile(r"^Episode\s+(\d+)\s*[-–—]?\s*(.*)$", re.IGNORECASE),
    re.compile(r"^Chapter\s+(\d+)\s*[-–—]?\s*(.*)$", re.IGNORECASE),
    re.compile(r"^Part\s+(\d+)\s*[-–—]?\s*(.*)$", re.IGNORECASE),
]


def detect_episode(text: str) -> Optional[dict]:
    """Return {number, title, heading} if text matches an episode heading."""
    text = text.strip()
    for pat in EPISODE_PATTERNS:
        m = pat.match(text)
        if m:
            return {
                "number": int(m.group(1)),
                "title": m.group(2).strip() if len(m.groups()) > 1 else "",
                "heading": text,
            }
    return None


# ─────────────────────────────────────────────────────────────
# DOCX analysis
# ─────────────────────────────────────────────────────────────
def analyse_docx(doc_path: str) -> dict:
    doc = Document(doc_path)
    paras = doc.paragraphs
    episodes: list[dict] = []
    current: Optional[dict] = None

    for i, para in enumerate(paras):
        ep = detect_episode(para.text)
        if ep:
            if current is not None:
                current["end_idx"] = i - 1
                episodes.append(current)
            current = {**ep, "start_idx": i, "end_idx": None, "word_count": 0}
        elif current is not None:
            current["word_count"] += len(para.text.split())

    if current is not None:
        current["end_idx"] = len(paras) - 1
        episodes.append(current)

    total_words = sum(e["word_count"] for e in episodes)
    stats = {
        "total_episodes": len(episodes),
        "total_words": total_words,
        "avg_words": round(total_words / len(episodes)) if episodes else 0,
        "largest": max(episodes, key=lambda e: e["word_count"], default=None),
        "smallest": min(episodes, key=lambda e: e["word_count"], default=None),
    }
    return {"episodes": episodes, "stats": stats}


# ─────────────────────────────────────────────────────────────
# DOCX writing
# ─────────────────────────────────────────────────────────────
def copy_paragraph(src_para, dest_doc):
    new_para = dest_doc.add_paragraph()
    try:
        sn = src_para.style.name
        if sn in [s.name for s in dest_doc.styles]:
            new_para.style = dest_doc.styles[sn]
    except Exception:
        pass
    try:
        if src_para.paragraph_format.alignment is not None:
            new_para.paragraph_format.alignment = src_para.paragraph_format.alignment
    except Exception:
        pass
    for run in src_para.runs:
        nr = new_para.add_run(run.text)
        try:
            nr.bold = run.bold
            nr.italic = run.italic
            nr.underline = run.underline
            if run.font.size:
                nr.font.size = run.font.size
            if run.font.name:
                nr.font.name = run.font.name
            if run.font.color and run.font.color.type:
                nr.font.color.rgb = run.font.color.rgb
        except Exception:
            pass
    return new_para


def build_docx(src_path: str, episodes: list[dict], ep_indices: list[int], out_path: str):
    src = Document(src_path)
    all_paras = src.paragraphs
    new_doc = Document()
    # Remove default blank paragraph
    for p in new_doc.paragraphs:
        p._element.getparent().remove(p._element)

    first = episodes[ep_indices[0]]["start_idx"]
    last = episodes[ep_indices[-1]]["end_idx"]
    for para in all_paras[first: last + 1]:
        copy_paragraph(para, new_doc)
    new_doc.save(out_path)


# ─────────────────────────────────────────────────────────────
# Pydantic models
# ─────────────────────────────────────────────────────────────
class SplitRequest(BaseModel):
    job_id: str
    mode: str                                  # "n_files" | "per_file" | "custom"
    n_files: Optional[int] = None
    per_file: Optional[int] = None
    custom_ranges: Optional[list[str]] = None  # e.g. ["1-20", "21-40"]


# ─────────────────────────────────────────────────────────────
# Routes
# ─────────────────────────────────────────────────────────────
@app.get("/health")
def health():
    return {"status": "ok", "version": "1.0.0"}


@app.post("/upload")
async def upload(file: UploadFile = File(...)):
    if not (file.filename or "").endswith(".docx"):
        raise HTTPException(400, "Only .docx files are accepted.")

    job_id = str(uuid.uuid4())
    job_dir = UPLOAD_DIR / job_id
    job_dir.mkdir()
    doc_path = str(job_dir / "source.docx")

    with open(doc_path, "wb") as f:
        f.write(await file.read())

    try:
        result = analyse_docx(doc_path)
    except Exception as exc:
        shutil.rmtree(job_dir, ignore_errors=True)
        raise HTTPException(422, f"Could not parse document: {exc}")

    if not result["episodes"]:
        shutil.rmtree(job_dir, ignore_errors=True)
        raise HTTPException(
            422,
            "No episodes found. Headings must match patterns like  'Ep 1 - Title'  or  'Chapter 1 - Title'.",
        )

    JOBS[job_id] = {
        "doc_path": doc_path,
        "job_dir": str(job_dir),
        "filename": file.filename,
        **result,
    }

    return {"job_id": job_id, "filename": file.filename, **result}


@app.post("/split")
def split(req: SplitRequest):
    job = JOBS.get(req.job_id)
    if not job:
        raise HTTPException(404, "Job not found — please upload again.")

    episodes = job["episodes"]
    n = len(episodes)
    doc_path = job["doc_path"]
    job_dir = Path(job["job_dir"])
    out_dir = job_dir / "output"
    out_dir.mkdir(exist_ok=True)

    # ── Compute groups ──────────────────────────────────────
    groups: list[tuple[int, int]] = []

    if req.mode == "n_files":
        k = max(1, req.n_files or 1)
        base, extra = divmod(n, k)
        idx = 0
        for i in range(k):
            size = base + (1 if i < extra else 0)
            if size:
                groups.append((idx, idx + size - 1))
                idx += size

    elif req.mode == "per_file":
        p = max(1, req.per_file or 1)
        for start in range(0, n, p):
            groups.append((start, min(start + p - 1, n - 1)))

    elif req.mode == "custom":
        if not req.custom_ranges:
            raise HTTPException(400, "custom_ranges required for custom mode.")
        ep_map = {e["number"]: i for i, e in enumerate(episodes)}
        for r in req.custom_ranges:
            parts = r.strip().split("-")
            if len(parts) != 2:
                raise HTTPException(400, f"Invalid range '{r}'. Use format '1-20'.")
            a, b = int(parts[0]), int(parts[1])
            if a not in ep_map or b not in ep_map:
                raise HTTPException(400, f"Episode numbers {a}–{b} not in document.")
            groups.append((ep_map[a], ep_map[b]))
    else:
        raise HTTPException(400, f"Unknown mode '{req.mode}'.")

    # ── Build DOCX files ─────────────────────────────────────
    output_files = []
    for g_start, g_end in groups:
        ep_s = episodes[g_start]["number"]
        ep_e = episodes[g_end]["number"]
        fname = f"Episodes_{ep_s:03d}_{ep_e:03d}.docx"
        out_path = str(out_dir / fname)
        build_docx(doc_path, episodes, list(range(g_start, g_end + 1)), out_path)
        output_files.append({"filename": fname, "ep_start": ep_s, "ep_end": ep_e, "path": out_path})

    # ── Pack ZIP ─────────────────────────────────────────────
    zip_path = str(job_dir / "Split_Document.zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in output_files:
            zf.write(f["path"], f["filename"])

    JOBS[req.job_id]["output_files"] = output_files
    JOBS[req.job_id]["zip_path"] = zip_path

    return {
        "job_id": req.job_id,
        "files": [{"filename": f["filename"], "ep_start": f["ep_start"], "ep_end": f["ep_end"]} for f in output_files],
        "zip_filename": "Split_Document.zip",
    }


@app.get("/download/{job_id}/zip")
def download_zip(job_id: str):
    job = JOBS.get(job_id)
    if not job:
        raise HTTPException(404, "Job not found.")
    zp = job.get("zip_path")
    if not zp or not os.path.exists(zp):
        raise HTTPException(404, "ZIP not ready — call /split first.")
    return FileResponse(zp, media_type="application/zip", filename="Split_Document.zip")


@app.get("/download/{job_id}/file/{filename}")
def download_file(job_id: str, filename: str):
    job = JOBS.get(job_id)
    if not job:
        raise HTTPException(404, "Job not found.")
    match = next((f for f in job.get("output_files", []) if f["filename"] == filename), None)
    if not match or not os.path.exists(match["path"]):
        raise HTTPException(404, "File not found.")
    return FileResponse(
        match["path"],
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
    )


if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=8765, reload=True, log_level="info")
